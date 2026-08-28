"""Gera artigo.docx a partir dos fontes LaTeX do artigo, para leitura fora do LaTeX.

Requer Pandoc, python-docx e, no workflow, LibreOffice. O Pandoc converte o
texto e as referências; o script reconstrói tabelas, corrige referências
cruzadas, insere fluxogramas TikZ e, por fim, abre e resalva o documento pelo
LibreOffice para maximizar a compatibilidade com Microsoft Word, LibreOffice e
visualizadores móveis.
"""
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
LATEX_DIR = ROOT / "latex-artigo"
INTERMEDIARIO = LATEX_DIR / "_main_pandoc_bruto.docx"
DESTINO = ROOT / "artigo.docx"

COLSPEC = re.compile(r"^(?:(?:>p[\d.]+cm|Y)\s*)+")

# O leitor LaTeX do Pandoc nao reconhece tabelas com tabularx/multicolumn como
# tabela nativa (cai no fallback de texto que construir_tabela() reconstroi) e,
# nesse modo, \toprule/\midrule/\bottomrule/\cmidrule viram texto cru (ex.:
# "(lr)1-6") e o conteudo de \multicolumn{N}{align}{texto} e descartado por
# completo (o rotulo da linha de subtotal desaparece). Por isso o texto e
# pre-processado (em copia temporaria, sem tocar os .tex versionados que o
# pipeline do PDF usa) antes de chegar ao Pandoc.
REGRA_TABELA_RE = re.compile(r"\\(?:top|mid|bottom|cmid)rule(?:\([a-z]{1,2}\))?(?:\{[^}]*\})?")
MULTICOLUNA_RE = re.compile(r"\\multicolumn\{(\d+)\}\{[^{}]*\}\{((?:[^{}]|\{[^{}]*\})*)\}")
# Marcador para as celulas "vazias" geradas ao desfazer um \multicolumn: o
# Pandoc colapsa sequencias de espacos em branco, entao duas celulas vazias
# separadas so por " & " (sem texto entre elas) acabam se fundindo num unico
# "&" solto na hora de dividir a linha por " & ". Um marcador nao-vazio evita
# essa colisao; construir_tabela() troca o marcador de volta por "".
CELULA_VAZIA = "@@CELULA_VAZIA@@"
# O codigo bruto do TikZ (com seus varios "&" de posicionamento de nos) tambem
# engana a heuristica de deteccao de tabela e vira uma tabela vazia no .docx.
# As figuras TikZ sao inseridas como imagem por inserir_figuras_tikz(), que le
# os .tex originais (nao esta copia temporaria) — remover o bloco aqui nao
# afeta a insercao da imagem, so evita a falsa tabela.
TIKZPICTURE_RE = re.compile(r"\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}", re.DOTALL)

# O Pandoc tambem descarta por completo o \caption{...} das 11 tabelas (elas
# nao viram tabela nativa, entao o float e todo o seu \caption somem, restando
# so o conteudo cru do tabularx) — o titulo da tabela desaparecia do .docx
# inteiro, nao so o rotulo do subtotal. Convertido aqui num paragrafo de texto
# comum antes do Pandoc rodar; corrigir_legendas_tabela() o transforma de
# volta num paragrafo com o estilo "Table Caption" e o numero correto.
LEGENDA_TABELA_RE = re.compile(r"\\caption\{((?:[^{}]|\{[^{}]*\})*)\}\s*\\label\{tab:([^}]+)\}")
MARCADOR_LEGENDA = "@@LEGENDA_TABELA@@"


def _marcar_legenda_tabela(casamento):
    texto, rotulo = casamento.group(1), casamento.group(2)
    return f"\n\n{MARCADOR_LEGENDA}tab:{rotulo}{MARCADOR_LEGENDA}{texto}{MARCADOR_LEGENDA}\n\n"


def _expandir_multicolumn(casamento):
    n = int(casamento.group(1))
    texto = casamento.group(2)
    if n > 1:
        return texto + (" & " + CELULA_VAZIA) * (n - 1)
    return texto


def preparar_fonte_pandoc():
    """Copia latex-artigo para um diretorio temporario com os comandos de
    regra/multicoluna das tabelas neutralizados, para o Pandoc ler."""
    temporario = Path(tempfile.mkdtemp(prefix="artigo_pandoc_fonte_"))
    destino = temporario / "latex-artigo"
    shutil.copytree(LATEX_DIR, destino)
    for arquivo in (destino / "sections").glob("*.tex"):
        texto = arquivo.read_text(encoding="utf-8")
        texto = REGRA_TABELA_RE.sub("", texto)
        texto = MULTICOLUNA_RE.sub(_expandir_multicolumn, texto)
        texto = TIKZPICTURE_RE.sub("", texto)
        texto = LEGENDA_TABELA_RE.sub(_marcar_legenda_tabela, texto)
        arquivo.write_text(texto, encoding="utf-8")
    return destino


def rodar_pandoc():
    pandoc = shutil.which("pandoc")
    if not pandoc:
        sys.exit("Pandoc nao encontrado no PATH.")
    fonte = preparar_fonte_pandoc()
    subprocess.run(
        [
            pandoc,
            "main.tex",
            "--bibliography=references.bib",
            "--csl=abnt.csl",
            "--citeproc",
            "--number-sections",
            "-o",
            str(INTERMEDIARIO),
            "--resource-path=.;sections;figuras",
        ],
        cwd=fonte,
        check=True,
    )
    shutil.rmtree(fonte.parent, ignore_errors=True)


def converter_figuras_pdf_no_docx():
    """Substitui mídias PDF por PNG e atualiza relacionamentos OOXML."""
    with tempfile.TemporaryDirectory(prefix="artigo_docx_") as temporario:
        raiz = Path(temporario)
        with zipfile.ZipFile(INTERMEDIARIO) as pacote:
            pacote.extractall(raiz)

        pasta_midias = raiz / "word" / "media"
        midias_pdf = sorted(pasta_midias.glob("*.pdf")) if pasta_midias.exists() else []
        if not midias_pdf:
            return

        pdftoppm = shutil.which("pdftoppm")
        if not pdftoppm:
            sys.exit("pdftoppm nao encontrado para converter figuras PDF.")

        substituicoes = {}
        for origem in midias_pdf:
            destino = origem.with_suffix(".png")
            subprocess.run(
                [pdftoppm, "-png", "-singlefile", "-r", "200", str(origem), str(destino.with_suffix(""))],
                check=True,
            )
            substituicoes[origem.name] = destino.name
            origem.unlink()

        for relacoes in raiz.rglob("*.rels"):
            conteudo = relacoes.read_text(encoding="utf-8")
            atualizado = conteudo
            for antigo, novo in substituicoes.items():
                atualizado = atualizado.replace(f"media/{antigo}", f"media/{novo}")
            if atualizado != conteudo:
                relacoes.write_text(atualizado, encoding="utf-8")

        tipos = raiz / "[Content_Types].xml"
        conteudo_tipos = tipos.read_text(encoding="utf-8")
        if 'Extension="png"' not in conteudo_tipos:
            conteudo_tipos = conteudo_tipos.replace(
                "</Types>", '<Default Extension="png" ContentType="image/png"/></Types>'
            )
        conteudo_tipos = re.sub(
            r'<Default Extension="pdf" ContentType="application/pdf"\s*/>',
            "",
            conteudo_tipos,
        )
        tipos.write_text(conteudo_tipos, encoding="utf-8")

        convertido = INTERMEDIARIO.with_name("_main_pandoc_figuras_convertidas.docx")
        with zipfile.ZipFile(convertido, "w", zipfile.ZIP_DEFLATED) as pacote:
            for arquivo in sorted(raiz.rglob("*")):
                if arquivo.is_file():
                    pacote.write(arquivo, arquivo.relative_to(raiz))
        convertido.replace(INTERMEDIARIO)
        print(f"Figuras PDF convertidas para PNG: {len(midias_pdf)}")


def eh_paragrafo_de_tabela(texto):
    if " & " not in texto:
        return False
    primeira_linha = texto.split("\n", 1)[0]
    return bool(COLSPEC.match(primeira_linha)) or texto.count(" & ") >= 2


def aplicar_bordas(tabela):
    from docx.oxml.ns import qn

    tbl_pr = tabela._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_pr.makeelement(
            qn(f"w:{lado}"),
            {qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "0", qn("w:color"): "999999"},
        )
        borders.append(el)
    tbl_pr.append(borders)


def aplicar_cabecalho_repetido_e_quebra(tabela):
    """Repete a linha de cabecalho (primeira linha) em toda pagina nova em que
    a tabela continua, e marca todas as linhas como indivisiveis (cantSplit)
    para que nenhuma linha - sobretudo as de subtotal - fique cortada ao meio
    ou orfa entre duas paginas."""
    from docx.oxml.ns import qn

    for i, linha in enumerate(tabela.rows):
        tr_pr = linha._tr.get_or_add_trPr()
        tr_pr.append(tr_pr.makeelement(qn("w:cantSplit"), {}))
        if i == 0:
            tr_pr.append(tr_pr.makeelement(qn("w:tblHeader"), {}))


def construir_tabela(doc, texto):
    from docx.shared import Pt

    texto = COLSPEC.sub("", texto, count=1).strip()
    linhas = [ln for ln in texto.split("\n") if ln.strip()]
    matriz = [
        ["" if c.strip() == CELULA_VAZIA else c.strip() for c in ln.split(" & ")]
        for ln in linhas
    ]
    ncols = max(len(r) for r in matriz)
    for linha in matriz:
        linha.extend([""] * (ncols - len(linha)))
    tabela = doc.add_table(rows=len(matriz), cols=ncols)
    aplicar_bordas(tabela)
    aplicar_cabecalho_repetido_e_quebra(tabela)
    for i, linha in enumerate(matriz):
        for j, valor in enumerate(linha):
            cell = tabela.cell(i, j)
            cell.text = valor
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.font.bold = True
    return tabela


ORDEM_SECOES = ("01_introducao", "02_metodo", "03_resultados_discussao", "04_consideracoes_finais")


def calcular_numeracao_tab_fig():
    numeros = {}
    contadores = {"tab": 0, "fig": 0}
    for fonte in sorted((LATEX_DIR / "sections").glob("*.tex")):
        conteudo = fonte.read_text(encoding="utf-8")
        for tipo, rotulo in re.findall(r"\\label\{(tab|fig):([^}]+)\}", conteudo):
            contadores[tipo] += 1
            numeros[f"{tipo}:{rotulo}"] = str(contadores[tipo])
    return numeros


def calcular_numeracao_secoes():
    """Replica a numeracao que `pandoc --number-sections` aplica aos
    cabecalhos (mesma ordem de \\input de main.tex), para resolver as
    remissoes Secao~\\ref{sec:x} do texto para o mesmo numero que aparece no
    cabecalho e no PDF (onde o LaTeX numera \\section/\\subsection por
    padrao)."""
    numeros = {}
    secao_re = re.compile(r"^\\section\{")
    sub_re = re.compile(r"^\\subsection\{")
    label_re = re.compile(r"^\\label\{sec:([^}]+)\}")
    contador_secao = 0
    contador_sub = 0
    numero_atual = None
    for nome in ORDEM_SECOES:
        conteudo = (LATEX_DIR / "sections" / f"{nome}.tex").read_text(encoding="utf-8")
        for linha in conteudo.splitlines():
            if secao_re.match(linha):
                contador_secao += 1
                contador_sub = 0
                numero_atual = str(contador_secao)
            elif sub_re.match(linha):
                contador_sub += 1
                numero_atual = f"{contador_secao}.{contador_sub}"
            else:
                casamento = label_re.match(linha)
                if casamento and numero_atual is not None:
                    numeros[f"sec:{casamento.group(1)}"] = numero_atual
    return numeros


def corrigir_referencias_cruzadas(doc):
    from docx.oxml.ns import qn

    numeros = calcular_numeracao_tab_fig()
    numeros.update(calcular_numeracao_secoes())

    padrao = re.compile(r"\[((?:tab|fig|sec):[^\]]+)\]")
    total = 0
    for no_texto in doc.element.body.iter(qn("w:t")):
        original = no_texto.text or ""
        atualizado, quantidade = padrao.subn(lambda achado: numeros.get(achado.group(1), achado.group(0)), original)
        if quantidade:
            no_texto.text = atualizado
            total += quantidade
    print(f"Referencias cruzadas numeradas: {total}")


def corrigir_legendas_tabela(doc):
    """Reconstroi, como paragrafo de estilo 'Table Caption', a legenda de
    cada uma das 11 tabelas a partir do marcador de texto plano deixado por
    LEGENDA_TABELA_RE (ver preparar_fonte_pandoc) - sem essa reconstrucao o
    titulo da tabela simplesmente some do .docx (o Pandoc descarta o
    \\caption{} das tabelas que nao vira tabela nativa)."""
    padrao = re.compile(
        re.escape(MARCADOR_LEGENDA) + r"(tab:[^@]+)" + re.escape(MARCADOR_LEGENDA)
        + r"(.*)" + re.escape(MARCADOR_LEGENDA),
        re.DOTALL,
    )
    numeros = calcular_numeracao_tab_fig()
    estilo_legenda = doc.styles["Table Caption"] if "Table Caption" in (s.name for s in doc.styles) else None
    total = 0
    for paragrafo in doc.paragraphs:
        casamento = padrao.search(paragrafo.text)
        if not casamento:
            continue
        rotulo, texto = casamento.group(1), casamento.group(2).strip()
        numero = numeros.get(rotulo, "?")
        for run in list(paragrafo.runs):
            run.text = ""
        texto_final = f"Tabela {numero}. {texto}"
        if paragrafo.runs:
            paragrafo.runs[0].text = texto_final
        else:
            paragrafo.add_run(texto_final)
        if estilo_legenda is not None:
            paragrafo.style = estilo_legenda
        paragrafo.paragraph_format.keep_with_next = True
        total += 1
    print(f"Legendas de tabela reconstruidas: {total}")


def _remover_tabela_fantasma_antes(paragrafo):
    """O Pandoc converte um \\begin{figure} sem imagem (caso do TikZ, cuja
    figura real e inserida a seguir como PNG) numa <w:tbl> vazia (0 linhas)
    logo antes da legenda. Remove esse residuo, se existir."""
    from docx.oxml.ns import qn

    anterior = paragrafo._p.getprevious()
    if anterior is not None and anterior.tag == qn("w:tbl") and not anterior.findall(qn("w:tr")):
        anterior.getparent().remove(anterior)
        return True
    return False


def inserir_figuras_tikz(doc):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    pdflatex = shutil.which("pdflatex")
    pdftoppm = shutil.which("pdftoppm")
    if not pdflatex or not pdftoppm:
        sys.exit("pdflatex e pdftoppm sao necessarios para preservar fluxogramas TikZ.")

    padrao = re.compile(
        r"\\begin\{figure\}.*?(\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}).*?\\caption\{([^{}]+)\}",
        re.DOTALL,
    )
    blocos = []
    for fonte in sorted((LATEX_DIR / "sections").glob("*.tex")):
        blocos.extend(padrao.findall(fonte.read_text(encoding="utf-8")))

    legendas = {p.text.strip(): p for p in doc.paragraphs if p.style.name == "Image Caption"}
    inseridas = 0
    with tempfile.TemporaryDirectory(prefix="artigo_tikz_") as temporario:
        raiz = Path(temporario)
        for indice, (tikz, legenda) in enumerate(blocos, start=1):
            paragrafo_legenda = legendas.get(legenda.strip())
            if paragrafo_legenda is None:
                continue
            fonte_tex = raiz / f"figura_tikz_{indice}.tex"
            fonte_tex.write_text(
                r"\documentclass[tikz,border=5pt]{standalone}" "\n"
                r"\usepackage[utf8]{inputenc}" "\n" r"\usepackage[T1]{fontenc}" "\n"
                r"\usetikzlibrary{arrows.meta,positioning}" "\n" r"\begin{document}" "\n"
                f"{tikz}\n" r"\end{document}" "\n",
                encoding="utf-8",
            )
            subprocess.run(
                [pdflatex, "-interaction=nonstopmode", "-halt-on-error", fonte_tex.name],
                cwd=raiz,
                check=True,
                stdout=subprocess.DEVNULL,
            )
            pdf = fonte_tex.with_suffix(".pdf")
            png = fonte_tex.with_suffix(".png")
            subprocess.run([pdftoppm, "-png", "-singlefile", "-r", "200", str(pdf), str(png.with_suffix(""))], check=True)
            paragrafo_imagem = doc.add_paragraph()
            paragrafo_imagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragrafo_imagem.add_run().add_picture(str(png), width=Inches(6.2))
            _remover_tabela_fantasma_antes(paragrafo_legenda)
            paragrafo_legenda._p.addprevious(paragrafo_imagem._p)
            inseridas += 1
    print(f"Fluxogramas TikZ inseridos: {inseridas}")


def corrigir_docx():
    from docx import Document

    doc = Document(INTERMEDIARIO)
    alvos = [p for p in doc.paragraphs if eh_paragrafo_de_tabela(p.text)]
    for paragrafo in alvos:
        tabela = construir_tabela(doc, paragrafo.text)
        paragrafo._p.addnext(tabela._tbl)
        legenda = doc.add_paragraph()
        tabela._tbl.addnext(legenda._p)
        paragrafo._p.getparent().remove(paragrafo._p)
    print(f"Tabelas reconstruidas: {len(alvos)}")

    corrigir_legendas_tabela(doc)

    estilo_h1 = next((p.style for p in doc.paragraphs if p.text.strip() == "Introdução"), None)
    for paragrafo in doc.paragraphs:
        if paragrafo.style.name == "Bibliography":
            titulo = paragrafo.insert_paragraph_before("Referências")
            if estilo_h1 is not None:
                titulo.style = estilo_h1
            break

    corrigir_referencias_cruzadas(doc)
    inserir_figuras_tikz(doc)
    doc.save(DESTINO)
    if not DESTINO.exists() or DESTINO.stat().st_size == 0:
        sys.exit("Falha ao gerar o Word.")
    print(f"Word gerado com {len(doc.paragraphs)} paragrafos e {len(doc.tables)} tabelas.")


def normalizar_com_libreoffice():
    """Abre e resalva o DOCX para corrigir incompatibilidades OOXML silenciosas."""
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        print("LibreOffice nao encontrado; normalizacao adicional ignorada.")
        return
    with tempfile.TemporaryDirectory(prefix="normalizar_docx_") as temporario:
        saida = Path(temporario)
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", str(saida), str(DESTINO)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            timeout=180,
        )
        normalizado = saida / DESTINO.name
        if not normalizado.exists() or normalizado.stat().st_size == 0:
            sys.exit("LibreOffice nao produziu o DOCX normalizado.")
        normalizado.replace(DESTINO)
    print("Word aberto e resalvo com sucesso pelo LibreOffice.")


def aplicar_formatacao_final():
    """Aplica A4/margens/numeracao de linha continua/Times New Roman 12,
    exigidos pela Ambiente Construido. Antes era um passo manual pos-hoc
    (scripts/python/ajustar_formato_word.py rodado a parte); agora roda
    sempre, depois do LibreOffice resalvar o arquivo (senao o proprio
    LibreOffice reverteria esses ajustes de secao/fonte)."""
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    import importlib

    ajustar = importlib.import_module("ajustar_formato_word")

    with zipfile.ZipFile(DESTINO) as zin:
        nomes = zin.namelist()
        conteudos = {n: zin.read(n) for n in nomes}

    doc_xml = conteudos["word/document.xml"].decode("utf-8")
    doc_xml, n_sect = ajustar.ajustar_sect_pr(doc_xml)
    doc_xml, n_fontes_doc = ajustar.ajustar_fonte_padrao(doc_xml)
    conteudos["word/document.xml"] = doc_xml.encode("utf-8")

    styles_xml = conteudos["word/styles.xml"].decode("utf-8")
    styles_xml, n_fontes_estilos = ajustar.ajustar_fonte_padrao(styles_xml)
    conteudos["word/styles.xml"] = styles_xml.encode("utf-8")

    with zipfile.ZipFile(DESTINO, "w", zipfile.ZIP_DEFLATED) as zout:
        for nome in nomes:
            zout.writestr(nome, conteudos[nome])

    print(
        f"Formato final aplicado: {n_sect} secao(oes) ajustada(s) para A4/margens/numeracao de linha, "
        f"{n_fontes_doc + n_fontes_estilos} fonte(s) trocada(s) para Times New Roman."
    )


def validar_docx():
    from docx import Document

    if not zipfile.is_zipfile(DESTINO):
        sys.exit("O arquivo Word gerado nao e um pacote DOCX valido.")
    with zipfile.ZipFile(DESTINO) as pacote:
        erro = pacote.testzip()
        if erro:
            sys.exit(f"Entrada corrompida no DOCX: {erro}")
        obrigatorios = {"[Content_Types].xml", "word/document.xml", "word/_rels/document.xml.rels"}
        ausentes = obrigatorios.difference(pacote.namelist())
        if ausentes:
            sys.exit(f"Partes obrigatorias ausentes no DOCX: {sorted(ausentes)}")
    documento = Document(DESTINO)
    if not documento.paragraphs:
        sys.exit("O Word foi aberto, mas nao contem paragrafos.")
    print(f"Word validado: {DESTINO.stat().st_size} bytes.")


if __name__ == "__main__":
    rodar_pandoc()
    converter_figuras_pdf_no_docx()
    corrigir_docx()
    normalizar_com_libreoffice()
    aplicar_formatacao_final()
    validar_docx()
    INTERMEDIARIO.unlink(missing_ok=True)
