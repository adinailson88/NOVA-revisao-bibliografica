"""Verifica o artigo.docx final contra os achados da auditoria de adequacao
a Ambiente Construido (branch submissao-ambiente-construido): estilo de
citacao ABNT, ausencia de residuo LaTeX nas tabelas, numeracao de secoes/
tabelas/figuras, contagem de palavras, formatacao A4/margens/fonte/linha, e
paridade minima de conteudo cientifico (numeros-ancora do corpus).

Uso: python scripts/python/verificar_artigo_word.py [caminho/para/artigo.docx]
Sem argumento, verifica artigo.docx na raiz do repositorio. Sai com codigo
de erro (AssertionError) na primeira falha, como os demais verificadores do
projeto.
"""
from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
LATEX_DIR = ROOT / "latex-artigo"
SECOES = LATEX_DIR / "sections"

LIMITE_PALAVRAS = 7000
MARGEM_MINIMA_PALAVRAS = 100
META_CONSERVADORA_PALAVRAS = 6880
TAMANHO_MAXIMO_BYTES = 5 * 1024 * 1024

NUMEROS_ANCORA = (
    "12.118", "9.542", "3.678", "137", "104", "121", "372", "109",
    "15 critérios", "seis dimensões", "43",
)

# Mesmos padroes ja usados por verificar_artigo.py (Etapa 13) para declarar
# como NAO realizado / trabalho futuro o rastreamento de citacoes e a busca
# estruturada de literatura cinzenta -- checados aqui tambem contra o TEXTO
# FINAL DO WORD (nao so o .tex fonte), para garantir que o Pandoc nao os
# tenha descartado na conversao.
PADROES_LIMITACOES = (
    r"rastreamento de citaç(ões|ão)",
    r"literatura cinzenta",
)


def _texto_docx(doc):
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                partes.append(celula.text)
    return "\n".join(partes)


def verificar(caminho_docx: Path):
    from docx import Document

    if not caminho_docx.exists():
        raise AssertionError(f"artigo.docx nao encontrado em {caminho_docx}")

    tamanho = caminho_docx.stat().st_size
    if tamanho == 0:
        raise AssertionError("artigo.docx esta vazio.")
    if tamanho > TAMANHO_MAXIMO_BYTES:
        raise AssertionError(f"artigo.docx tem {tamanho} bytes, acima do limite de 5MB da revista.")

    if not zipfile.is_zipfile(caminho_docx):
        raise AssertionError("artigo.docx nao e um pacote ZIP/DOCX valido.")
    with zipfile.ZipFile(caminho_docx) as pacote:
        erro = pacote.testzip()
        if erro:
            raise AssertionError(f"Entrada corrompida no DOCX: {erro}")
        nomes = set(pacote.namelist())
        obrigatorios = {"[Content_Types].xml", "word/document.xml", "word/_rels/document.xml.rels"}
        ausentes = obrigatorios.difference(nomes)
        if ausentes:
            raise AssertionError(f"Partes obrigatorias ausentes no DOCX: {sorted(ausentes)}")
        doc_xml = pacote.read("word/document.xml").decode("utf-8")

    doc = Document(caminho_docx)
    texto_completo = _texto_docx(doc)

    # 1) "seis dimensões" e "15 critérios" na prosa fonte (.tex), nao so na tabela.
    texto_fonte = "\n".join(
        (SECOES / f"{nome}.tex").read_text(encoding="utf-8")
        for nome in ("00_resumo", "01_introducao", "02_metodo", "03_resultados_discussao", "04_consideracoes_finais")
    )
    if "seis dimensões" not in texto_fonte:
        raise AssertionError('Frase "seis dimensões" ausente da prosa do artigo.')
    if "15 critérios" not in texto_fonte:
        raise AssertionError('Frase "15 critérios" ausente da prosa do artigo.')

    # 2) Ausencia de residuo LaTeX de tabela no Word.
    for residuo in (r"\cmidrule", r"\toprule", r"\midrule", r"\bottomrule", "(lr)1-6", r"\multicolumn"):
        if residuo in texto_completo:
            raise AssertionError(f"Residuo LaTeX encontrado no Word: {residuo!r}")

    # 3) Nenhuma remissao ou rotulo LaTeX cru sobrando (\ref, \label, [tab:/fig:/sec:...]).
    if re.search(r"\[(?:tab|fig|sec):[^\]]+\]", texto_completo):
        raise AssertionError("Remissao cruzada nao resolvida ([tab:/fig:/sec:...]) encontrada no Word.")
    if "\\ref{" in texto_completo or "\\label{" in texto_completo:
        raise AssertionError("Comando \\ref{} ou \\label{} cru encontrado no Word.")

    # 4) Ausencia de " and " em citacao (padrao autor-data entre parenteses).
    for casamento in re.finditer(r"\([^()]*\)", texto_completo):
        trecho = casamento.group(0)
        if " and " in trecho and re.search(r",\s*\d{4}", trecho):
            raise AssertionError(f'Citação em estilo inglês (" and ") encontrada: {trecho[:80]!r}')

    # 5) 44 referencias na bibliografia (43 do corpus + o dataset no Zenodo).
    biblio = [p for p in doc.paragraphs if p.style.name == "Bibliography"]
    if len(biblio) != 44:
        raise AssertionError(f"Esperava 44 referências na bibliografia do Word, encontrou {len(biblio)}.")

    # 6) 11 tabelas e 11 figuras/graficos.
    if len(doc.tables) != 11:
        raise AssertionError(f"Esperava 11 tabelas no Word, encontrou {len(doc.tables)}.")
    legendas_tabela = [p for p in doc.paragraphs if p.style.name == "Table Caption"]
    if len(legendas_tabela) != 11:
        raise AssertionError(f"Esperava 11 legendas de tabela no Word, encontrou {len(legendas_tabela)}.")
    legendas_figura = [p for p in doc.paragraphs if p.style.name == "Image Caption"]
    if len(legendas_figura) != 11:
        raise AssertionError(f"Esperava 11 legendas de figura/gráfico no Word, encontrou {len(legendas_figura)}.")
    for legenda in legendas_figura:
        if re.match(r"^Gráfico\s+\.", legenda.text):
            raise AssertionError(f"Legenda de gráfico sem número: {legenda.text!r}")

    # 7) Numeracao de secoes: os titulos Heading 1/Heading 2 devem comecar com
    # numero (Pandoc --number-sections), e nao deve haver titulo "orfao" sem
    # numero exceto o proprio "Referências" (inserido a parte, sem numero).
    for paragrafo in doc.paragraphs:
        if paragrafo.style.name in ("Heading 1", "Heading 2") and paragrafo.text.strip() != "Referências":
            if not re.match(r"^\d+(\.\d+)*[\s\t]", paragrafo.text.strip() + " "):
                raise AssertionError(f"Título de seção sem numeração: {paragrafo.text!r}")

    # 8) Formatacao A4 / margens / fonte / numeracao de linha continua.
    secao = doc.sections[0]
    if abs(secao.page_width.cm - 21.0) > 0.1 or abs(secao.page_height.cm - 29.7) > 0.1:
        raise AssertionError(
            f"Página não está em A4: {secao.page_width.cm:.2f}x{secao.page_height.cm:.2f} cm."
        )
    margens_esperadas = {"top": 3.0, "left": 3.0, "bottom": 2.0, "right": 2.0}
    margens_reais = {
        "top": secao.top_margin.cm, "left": secao.left_margin.cm,
        "bottom": secao.bottom_margin.cm, "right": secao.right_margin.cm,
    }
    for lado, esperado in margens_esperadas.items():
        if abs(margens_reais[lado] - esperado) > 0.05:
            raise AssertionError(f"Margem {lado} é {margens_reais[lado]:.2f}cm, esperado {esperado}cm.")
    if "lnNumType" not in doc_xml:
        raise AssertionError("Numeração contínua de linhas (lnNumType) ausente no Word.")
    if "Times New Roman" not in doc_xml and "Times New Roman" not in zipfile.ZipFile(caminho_docx).read("word/styles.xml").decode("utf-8"):
        raise AssertionError("Fonte Times New Roman não encontrada no Word (nem em document.xml nem em styles.xml).")

    # 9) Contagem de palavras.
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    import importlib
    contador = importlib.import_module("contar_palavras_artigo")
    palavras = contador.contar_palavras(caminho_docx)
    if palavras >= LIMITE_PALAVRAS:
        raise AssertionError(f"{palavras} palavras >= limite de {LIMITE_PALAVRAS} da revista.")
    if LIMITE_PALAVRAS - palavras < MARGEM_MINIMA_PALAVRAS:
        raise AssertionError(
            f"Margem de apenas {LIMITE_PALAVRAS - palavras} palavras frente ao limite "
            f"(mínimo exigido: {MARGEM_MINIMA_PALAVRAS})."
        )
    if palavras > META_CONSERVADORA_PALAVRAS:
        raise AssertionError(
            f"{palavras} palavras acima da meta conservadora de {META_CONSERVADORA_PALAVRAS}."
        )

    # 10) Limitacoes declaradas como nao realizadas/trabalho futuro, tambem
    # presentes no texto final do Word (nao so no .tex).
    for padrao in PADROES_LIMITACOES:
        if not re.search(padrao, texto_completo):
            raise AssertionError(f"Limitação obrigatória ausente no Word: {padrao}")

    # 11) Declaração de uso de IA presente e com as ressalvas de escopo.
    if "Claude" not in texto_completo:
        raise AssertionError("Declaração de uso de IA (menção a Claude) ausente do Word.")
    if "não foi utilizada" not in texto_completo and "sem uso" not in texto_completo:
        raise AssertionError("Ressalva de escopo da declaração de uso de IA ausente do Word.")

    # 12) Paridade mínima de conteúdo científico: números-âncora do corpus
    # ainda presentes no Word (detecta perda real de conteúdo entre edições).
    for ancora in NUMEROS_ANCORA:
        if ancora not in texto_completo:
            raise AssertionError(f"Número/frase-âncora do corpus ausente do Word: {ancora!r}")

    print(
        f"Word verificado: {tamanho} bytes, {len(doc.tables)} tabelas, {len(legendas_figura)} figuras, "
        f"{len(biblio)} referências, {palavras} palavras (margem de {LIMITE_PALAVRAS - palavras} "
        f"frente ao limite de {LIMITE_PALAVRAS}). OK"
    )


if __name__ == "__main__":
    caminho = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "artigo.docx"
    verificar(caminho)
